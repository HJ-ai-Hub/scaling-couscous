#!/usr/bin/env python3
"""
Generates a professionally formatted Conditional Share Transfer Agreement (DOCX)
for the transfer of 21% of Orizuru Holiday Homes LLC from Podocarpus Real Estate LLC
to the two minority shareholders, governed by Dubai (UAE) law.

Output: agreements/Orizuru_Share_Transfer_Agreement.docx
PDF is produced separately via LibreOffice headless conversion.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Palette
# ----------------------------------------------------------------------------
NAVY = RGBColor(0x0E, 0x2A, 0x47)
GOLD = RGBColor(0xB1, 0x8A, 0x3E)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x8A, 0x8A, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

BODY_FONT = "Georgia"
HEAD_FONT = "Calibri"

doc = Document()

# ----------------------------------------------------------------------------
# Base styles
# ----------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = BLACK
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.18
pf.space_after = Pt(6)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_run(p, text, size=10.5, bold=False, italic=False, color=BLACK,
            font=BODY_FONT, caps=False):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    return r


def heading(num, text):
    """Top-level clause heading e.g. '1.  DEFINITIONS'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, f"{num}.", size=11.5, bold=True, color=NAVY, font=HEAD_FONT)
    add_run(p, "   ", size=11.5, bold=True, color=NAVY, font=HEAD_FONT)
    add_run(p, text.upper(), size=11.5, bold=True, color=NAVY, font=HEAD_FONT)
    return p


def clause(num, text, level=0):
    """Numbered sub-clause with hanging indent."""
    p = doc.add_paragraph()
    indent = 0.55 + level * 0.45
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.55)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, f"{num}", size=10.5, bold=True, color=NAVY)
    add_run(p, "\t", size=10.5)
    # allow simple **bold** segments
    parts = text.split("**")
    for i, seg in enumerate(parts):
        add_run(p, seg, size=10.5, bold=(i % 2 == 1))
    return p


def plain(text, justify=True, italic=False, color=BLACK, size=10.5,
          space_after=6, indent=0.0):
    p = doc.add_paragraph()
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    parts = text.split("**")
    for i, seg in enumerate(parts):
        add_run(p, seg, size=size, bold=(i % 2 == 1), italic=italic, color=color)
    return p


def recital(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.first_line_indent = Inches(-0.55)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, f"({label})", size=10.5, bold=True, color=NAVY)
    add_run(p, "\t", size=10.5)
    parts = text.split("**")
    for i, seg in enumerate(parts):
        add_run(p, seg, size=10.5, bold=(i % 2 == 1))
    return p


# ----------------------------------------------------------------------------
# Page geometry + header/footer
# ----------------------------------------------------------------------------
sec = doc.sections[0]
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.05)
sec.right_margin = Inches(1.05)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = HEAD_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT


# Footer
footer = sec.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(fp)
add_run(fp, "STRICTLY PRIVATE & CONFIDENTIAL", size=7.5, color=LIGHT,
        font=HEAD_FONT, caps=True)
fp2 = footer.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(fp2)
add_run(fp2, "Share Transfer Agreement — Orizuru Holiday Homes LLC   |   Page ",
        size=7.5, color=LIGHT, font=HEAD_FONT)
add_page_number_field(fp2)

# ----------------------------------------------------------------------------
# COVER PAGE
# ----------------------------------------------------------------------------
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "STRICTLY PRIVATE & CONFIDENTIAL", size=10, bold=True, color=GOLD,
        font=HEAD_FONT, caps=True)

# Gold rule
def rule(color="B18A3E", size="18"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

rule()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "SHARE TRANSFER AGREEMENT", size=30, bold=True, color=NAVY,
        font=HEAD_FONT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
add_run(p, "Conditional & Performance-Based", size=14, italic=True, color=GOLD,
        font=HEAD_FONT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "in respect of 21% of the issued share capital of",
        size=11, color=GREY, font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "ORIZURU HOLIDAY HOMES LLC", size=15, bold=True, color=NAVY,
        font=HEAD_FONT)

doc.add_paragraph()
rule(color="0E2A47", size="6")
doc.add_paragraph()

# Parties block on cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "BETWEEN", size=11, bold=True, color=GOLD, font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
add_run(p, "PODOCARPUS REAL ESTATE LLC", size=12.5, bold=True, color=NAVY,
        font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "(the Transferor)", size=10, italic=True, color=GREY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "AND", size=11, bold=True, color=GOLD, font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
add_run(p, "ORIZURU HOLIDAY HOMES LLC", size=12.5, bold=True, color=NAVY,
        font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "(the Company)", size=10, italic=True, color=GREY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "JOINED BY", size=10, bold=True, color=GOLD, font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
add_run(p, "MR. YSLAM ATAYEV   &   MR. ORAZDURDY MAMIYEV", size=11, bold=True,
        color=NAVY, font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "(together, the Transferees)", size=10, italic=True, color=GREY)

doc.add_paragraph()
doc.add_paragraph()
rule()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Dated: ____ / ____ / 20____", size=11, bold=True, color=NAVY,
        font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Governed by the laws of the United Arab Emirates as applied in the Emirate of Dubai",
        size=9.5, italic=True, color=GREY)

doc.add_page_break()

# ----------------------------------------------------------------------------
# IMPORTANT NOTICE / DISCLAIMER
# ----------------------------------------------------------------------------
def callout(title, body_lines, fill="F3EDDE", border="B18A3E", title_color=NAVY):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    # border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "10")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), border)
        borders.append(e)
    tcPr.append(borders)
    cell.paragraphs[0].text = ""
    tp = cell.paragraphs[0]
    no_space(tp)
    add_run(tp, title, size=10, bold=True, color=title_color, font=HEAD_FONT, caps=True)
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(4)
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parts = line.split("**")
        for i, seg in enumerate(parts):
            add_run(bp, seg, size=9, bold=(i % 2 == 1), color=GREY)
    # padding
    pad = cell.add_paragraph()
    no_space(pad)
    add_run(pad, "", size=3)
    return tbl

# remove the empty spacer heading we added
# (it created an empty paragraph; acceptable)

callout(
    "Important Legal Notice",
    [
        "This document is a template share transfer agreement prepared for the parties named above. It is drafted to reflect the principles of **UAE Federal Decree-Law No. (32) of 2021 on Commercial Companies** and the practice of the **Department of Economy and Tourism (Dubai)**. It is **not** a substitute for independent legal advice.",
        "A transfer of shares (an interest / quota) in a Dubai mainland Limited Liability Company is only fully effective once it is **notarised before a Notary Public in Dubai**, the **Memorandum of Association (MOA) is amended**, and the change is **registered with the Department of Economy and Tourism**. The parties should engage a UAE-licensed legal practitioner to review, finalise, prepare the Arabic notarised addendum and complete registration before execution.",
        "Fields shown as **[ ● ]** or blank lines are to be completed by the parties. A consolidated list of items to complete appears at **Schedule 5**.",
    ],
)

doc.add_paragraph()

# ----------------------------------------------------------------------------
# TITLE + DATE LINE (body)
# ----------------------------------------------------------------------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "SHARE TRANSFER AGREEMENT", size=16, bold=True, color=NAVY, font=HEAD_FONT)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
add_run(p, "(Conditional & Performance-Based)", size=11, italic=True, color=GOLD)

plain("This Share Transfer Agreement (this “**Agreement**”) is made on this "
      "______ day of ____________________ 20____ (the “**Effective Date**”).",
      space_after=10)

# ----------------------------------------------------------------------------
# PARTIES
# ----------------------------------------------------------------------------
def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    add_run(p, text.upper(), size=12, bold=True, color=NAVY, font=HEAD_FONT)
    rline = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "B18A3E")
    pbdr.append(bottom)
    rline.append(pbdr)
    return p

# remove the empty heading("","") artifacts by writing proper content now
section_title("Parties")

clause("(1)", "**PODOCARPUS REAL ESTATE LLC**, a limited liability company incorporated and "
       "licensed in the Emirate of Dubai, United Arab Emirates under commercial licence "
       "no. [ ● ], having its registered office at [ ● ], Dubai, UAE "
       "(the “**Transferor**” or “**Podocarpus**”);")
clause("(2)", "**ORIZURU HOLIDAY HOMES LLC**, a limited liability company incorporated and "
       "licensed in the Emirate of Dubai, United Arab Emirates under commercial licence "
       "no. [ ● ], having its registered office at [ ● ], Dubai, UAE "
       "(the “**Company**”);")
clause("(3)", "**MR. YSLAM ATAYEV**, holder of passport / Emirates ID no. [ ● ], "
       "of [ ● ] (“**Mr. Atayev**”); and")
clause("(4)", "**MR. ORAZDURDY MAMIYEV**, holder of passport / Emirates ID no. [ ● ], "
       "of [ ● ] (“**Mr. Mamiyev**”).")

plain("Mr. Atayev and Mr. Mamiyev are together referred to as the “**Transferees**” "
      "and each a “**Transferee**”. The Transferor, the Company and the Transferees "
      "are together referred to as the “**Parties**” and each a “**Party**”. "
      "The Transferor and the Company are the principal counterparties to this Agreement, and "
      "the Transferees join in and execute this Agreement as the recipients of the Transfer "
      "Shares and as continuing shareholders of the Company.")

# ----------------------------------------------------------------------------
# RECITALS
# ----------------------------------------------------------------------------
section_title("Recitals (Whereas)")

recital("A", "The Company is a limited liability company duly established in the Emirate of "
        "Dubai and engaged, among other activities, in the business of holiday homes / "
        "short-term rental and the leasing and operation of residential property in Dubai.")
recital("B", "As at the Effective Date, the issued share capital of the Company is held as "
        "follows: **Podocarpus – 70%**; **Mr. Atayev – 16%**; and "
        "**Mr. Mamiyev – 14%** (the “**Existing Shareholding**”), as more "
        "particularly set out in **Schedule 1**.")
recital("C", "The paid-up capital initially contributed to and deployed in the Company "
        "attributable to the Transferor is **AED 200,000 (two hundred thousand UAE Dirhams)** "
        "(the “**Initial Capital**”).")
recital("D", "The Transferor wishes to incentivise and reward the continuing efforts and "
        "performance of the Transferees in growing the business of the Company by agreeing to "
        "transfer to them, conditionally and as a performance reward, an aggregate of "
        "**21% (twenty-one per cent)** of the issued share capital of the Company, subject to "
        "and conditional upon the achievement of the Performance Threshold described in this "
        "Agreement.")
recital("E", "It is the express intention of the Parties that (i) the right to effect the "
        "transfer, and the determination of how the 21% is allocated between the Transferees "
        "according to their respective performance, shall vest in and be exercised by the "
        "Founder of the Transferor in the Founder’s **absolute and sole discretion**; "
        "(ii) no transfer shall occur unless and until the Performance Threshold is satisfied; "
        "and (iii) notwithstanding that the Transferor’s holding will reduce to 49% and the "
        "Transferees’ combined holding will become 51% upon Completion, the Transferor "
        "(acting by the Founder) shall **retain absolute decision-making power and control over "
        "the Company** following Completion, on the terms set out in Clause 7.")
recital("F", "The Parties enter into this Agreement to record the binding terms upon which the "
        "said transfer shall, upon satisfaction of the conditions set out herein, be carried "
        "into effect.")

# ----------------------------------------------------------------------------
# OPERATIVE
# ----------------------------------------------------------------------------
section_title("Operative Provisions")
plain("**NOW IT IS HEREBY AGREED** as follows:", space_after=8)

# 1. Definitions
heading("1", "Definitions and Interpretation")
clause("1.1", "In this Agreement, unless the context otherwise requires, the following "
       "expressions have the following meanings:", level=0)

defs = [
    ("AED", "the lawful currency of the United Arab Emirates (UAE Dirham)."),
    ("Applicable Law", "the federal laws of the United Arab Emirates and the laws, regulations "
     "and administrative requirements applicable in the Emirate of Dubai, including UAE Federal "
     "Decree-Law No. (32) of 2021 on Commercial Companies (as amended) (the “CCL”) and "
     "the requirements of the Department of Economy and Tourism, Dubai (“DET”)."),
    ("Allocation Notice", "the written notice issued by the Founder pursuant to Clause 4 "
     "determining the respective percentages of the Transfer Shares to be transferred to each "
     "Transferee, substantially in the form of Schedule 3."),
    ("Capital Recovery Amount", "AED 200,000 (two hundred thousand UAE Dirhams), being the amount "
     "of Cumulative Gross Profit notionally attributable to the full recovery of the Initial "
     "Capital."),
    ("Completion", "completion of the transfer of the Transfer Shares in accordance with "
     "Clause 6, including notarisation and registration with the DET."),
    ("Cumulative Gross Profit", "the aggregate gross profit of the Company (being total revenue "
     "actually realised and received by the Company less the direct cost of generating that "
     "revenue) accumulated from the commencement of the Company’s operations up to the date "
     "of measurement, as determined in accordance with Schedule 2."),
    ("Encumbrance", "any mortgage, charge, pledge, lien, option, restriction, right of first "
     "refusal, right of pre-emption, third-party right or interest, or any other encumbrance or "
     "security interest of any kind."),
    ("Founder", "[ ●  full name ], holder of passport / Emirates ID no. [ ● ], being "
     "the founder of the Transferor, or such other person as the Transferor may by written "
     "notice designate as exercising the Founder’s rights under this Agreement."),
    ("Performance Threshold", "the achievement by the Company of Cumulative Gross Profit of not "
     "less than AED 400,000 (four hundred thousand UAE Dirhams), comprising the Capital Recovery "
     "Amount of AED 200,000 plus the Surplus Profit Amount of AED 200,000, as more particularly "
     "described in Clause 3 and Schedule 2."),
    ("Surplus Profit Amount", "AED 200,000 (two hundred thousand UAE Dirhams), being the "
     "additional Cumulative Gross Profit required to be generated for the Company over and above "
     "the Capital Recovery Amount."),
    ("Transfer Shares", "an aggregate of 21% (twenty-one per cent) of the issued share capital "
     "of the Company currently held by the Transferor, to be transferred to the Transferees in "
     "the proportions determined under the Allocation Notice."),
    ("Trigger Event", "satisfaction of the Performance Threshold, being the event upon which the "
     "Transferor’s obligation to transfer the Transfer Shares (subject to the Founder’s "
     "rights under Clause 5) arises."),
]
for term, meaning in defs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.0)
    p.paragraph_format.first_line_indent = Inches(-0.45)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, f"“{term}”", size=10.5, bold=True, color=NAVY)
    add_run(p, f"  means {meaning}", size=10.5)

clause("1.2", "**Interpretation.** Clause headings are for convenience only and do not affect "
       "interpretation. References to a “Clause” or “Schedule” are to clauses "
       "of and schedules to this Agreement; the Schedules form part of this Agreement. Words "
       "importing the singular include the plural and vice versa. References to a “person” "
       "include natural persons and bodies corporate. A reference to a statute or statutory "
       "provision is a reference to it as amended, consolidated or re-enacted from time to time. "
       "“Including” means including without limitation.")

# 2. Subject Matter
heading("2", "Subject Matter and Conditional Nature")
clause("2.1", "This Agreement records the binding terms upon which the Transferor agrees to "
       "transfer, and the Transferees agree to acquire, the Transfer Shares as a performance "
       "reward, **conditional upon and subject to** the occurrence of the Trigger Event and the "
       "issuance of the Allocation Notice.")
clause("2.2", "The Parties acknowledge and agree that, notwithstanding any other provision of "
       "this Agreement, **no legal or beneficial interest in the Transfer Shares shall pass** to "
       "either Transferee unless and until (a) the Trigger Event has occurred; (b) the Founder "
       "has issued the Allocation Notice; and (c) Completion has taken place in accordance with "
       "Clause 6.")
clause("2.3", "Until Completion, the Transferor shall remain the legal and beneficial owner of "
       "the entirety of its 70% shareholding and shall be entitled to exercise all rights "
       "attaching to those shares, including voting and dividend rights.")

# 3. Performance Threshold / Trigger
heading("3", "Performance Threshold (the Sales / Profit Condition)")
clause("3.1", "The transfer of the Transfer Shares is conditional upon the Company achieving the "
       "**Performance Threshold**, namely Cumulative Gross Profit of not less than "
       "**AED 400,000 (four hundred thousand UAE Dirhams)**.")
clause("3.2", "The Performance Threshold shall be deemed to comprise two components: "
       "(a) the **Capital Recovery Amount** of AED 200,000, representing the full recovery to the "
       "Company of the Initial Capital of AED 200,000; and (b) the **Surplus Profit Amount** of "
       "AED 200,000, representing additional profit generated for the Company over and above the "
       "Capital Recovery Amount.")
clause("3.3", "Cumulative Gross Profit shall be measured on a cumulative basis from the "
       "commencement of the Company’s operations and shall be determined in accordance with "
       "the methodology set out in **Schedule 2**. There shall be **no fixed deadline** by which "
       "the Performance Threshold must be achieved; the condition is a continuing one and the "
       "Trigger Event occurs as and when the Performance Threshold is first satisfied.")
clause("3.4", "The Trigger Event shall be evidenced by management accounts of the Company "
       "certified by the Company’s accountant or auditor confirming that the Performance "
       "Threshold has been met (the “**Threshold Certificate**”). The Transferor (acting "
       "by the Founder) shall be entitled, but not obliged, to require independent verification "
       "of the Threshold Certificate at the Company’s cost before acting upon it.")
clause("3.5", "For the avoidance of doubt, the achievement of the Performance Threshold is a "
       "condition to the Transferor’s obligation to transfer; it does not, of itself, effect "
       "any transfer, which remains subject to the Founder’s rights under Clause 5 and to "
       "Completion under Clause 6.")

# 4. Allocation
heading("4", "Allocation of the Transfer Shares Between the Transferees")
clause("4.1", "The 21% Transfer Shares shall be allocated between the Transferees **according to "
       "their respective performance** and contribution to the achievement of the Performance "
       "Threshold and to the growth of the Company’s business.")
clause("4.2", "The determination of the respective percentages allocated to each Transferee "
       "shall be made by the **Founder in the Founder’s absolute and sole discretion**, and "
       "shall be recorded in the Allocation Notice issued substantially in the form of "
       "**Schedule 3**. The aggregate of the percentages so allocated shall equal, but shall not "
       "exceed, 21%.")
clause("4.3", "The Founder’s determination in the Allocation Notice shall, save in the case "
       "of manifest error or fraud, be **final and binding** on the Company and the Transferees, "
       "who irrevocably agree to accept and give effect to it.")
clause("4.4", "By way of illustration only and without binding the Founder, examples of possible "
       "resulting shareholding outcomes are set out in **Schedule 4**. Nothing in Schedule 4 "
       "fetters the Founder’s discretion under this Clause 4.")

# 5. Founder's reserved rights / absolute discretion
heading("5", "Founder’s Absolute Rights and Discretion")
clause("5.1", "The Parties expressly acknowledge and agree that the right to trigger, effect, "
       "structure, time, allocate, defer or decline the transfer of the Transfer Shares is "
       "vested in the Founder of the Transferor and is exercisable in the Founder’s "
       "**absolute and sole discretion**.")
clause("5.2", "Without limiting Clause 5.1, the Founder shall have the absolute right to "
       "determine: (a) whether and when the Trigger Event has been satisfied; (b) the respective "
       "allocation of the Transfer Shares between the Transferees; (c) whether any conditions of "
       "good standing, continued involvement or conduct of a Transferee have been met; and "
       "(d) the timing of Completion.")
clause("5.3", "The exercise (or non-exercise) by the Founder of any discretion under this "
       "Agreement shall not give rise to any claim, liability or cause of action against the "
       "Founder or the Transferor, provided such discretion is exercised in good faith and not "
       "fraudulently.")
clause("5.4", "The rights reserved to the Founder under this Clause 5 are in addition to, and not "
       "in substitution for, any rights of the Transferor as a shareholder under the "
       "Company’s MOA or Applicable Law.")

# 6. Completion
heading("6", "Completion, Notarisation and Registration")
clause("6.1", "Within 30 (thirty) business days following the later of (a) the occurrence of the "
       "Trigger Event and (b) the issuance of the Allocation Notice, the Parties shall complete "
       "the transfer of the Transfer Shares (“Completion”).")
clause("6.2", "At Completion, the Parties shall do all such acts and execute all such documents "
       "as are necessary under Applicable Law to give effect to the transfer, including: "
       "(a) executing an addendum to the Company’s Memorandum of Association reflecting the "
       "revised shareholding; (b) attending before a **Notary Public in Dubai** to notarise the "
       "share transfer and the MOA addendum; (c) procuring the issuance of an amended trade "
       "licence and the registration of the revised shareholding with the **DET**; and "
       "(d) updating the Company’s register of shareholders.")
clause("6.3", "The Company and each shareholder shall provide all consents, approvals, board and "
       "shareholder resolutions, powers of attorney and signatures reasonably required to effect "
       "Completion, and irrevocably undertake not to obstruct or delay Completion once the "
       "conditions in Clause 6.1 are satisfied.")
clause("6.4", "If any governmental, regulatory or DET approval is required for the transfer, the "
       "Parties shall co-operate in good faith to obtain it as soon as reasonably practicable.")

# 7. Post-Completion Governance — Reserved Control
heading("7", "Post-Completion Governance: Reserved Control of the Transferor")
clause("7.1", "**Fundamental Condition; Retention of Absolute Decision-Making.** The Parties "
       "acknowledge and agree that, although the Transferor’s shareholding will reduce from "
       "70% to 49% upon Completion and the Transferees’ combined holding will increase to 51%, "
       "it is the express and fundamental intention of the Parties that the Transferor "
       "(acting by the Founder) shall **retain absolute decision-making power and control over "
       "the Company** at all times following Completion. This is a fundamental basis upon which "
       "the Transferor has agreed to transfer the Transfer Shares, and the Transferees would "
       "not have been offered, and shall not be entitled to receive or to retain, the Transfer "
       "Shares but for the arrangements set out in this Clause 7.")
clause("7.2", "**Irrevocable Voting Undertaking.** Each Transferee irrevocably undertakes and "
       "covenants to the Transferor that, in respect of every matter put to a vote of the "
       "shareholders of the Company (whether at an ordinary general assembly, an extraordinary "
       "general assembly, by written resolution or otherwise) and in respect of all consents, "
       "approvals and decisions of the shareholders, he shall: (a) vote (and procure that any "
       "proxy or representative on his behalf shall vote) all of the shares held by him "
       "strictly in accordance with the written instructions of the Transferor (acting by the "
       "Founder); and (b) abstain from voting where so instructed in writing by the Transferor.")
clause("7.3", "**Irrevocable Proxy and Power of Attorney.** As security for the undertaking in "
       "Clause 7.2, each Transferee shall, at Completion and from time to time thereafter at "
       "the request of the Transferor, execute an irrevocable proxy and power of attorney in "
       "favour of the Transferor (or its nominee) authorising the Transferor to attend, speak "
       "and vote at all meetings of the shareholders of the Company in respect of the "
       "Transferee’s shares, and to execute on the Transferee’s behalf any written "
       "resolution or shareholder consent. To the maximum extent permitted by Applicable Law, "
       "such proxy and power of attorney shall be **irrevocable and coupled with an interest**.")
clause("7.4", "**Casting / Deciding Vote.** Where, in respect of any matter requiring the consent "
       "or decision of the shareholders of the Company, the view of the Transferor differs from "
       "that of either or both of the Transferees, the **view of the Transferor shall prevail "
       "and shall constitute the binding decision of the shareholders**. The Parties shall "
       "procure that the amended MOA expressly confers upon the Transferor a casting / deciding "
       "vote to give effect to this Clause.")
clause("7.5", "**Sole Right to Appoint the General Manager.** The Transferor shall have the sole "
       "and exclusive right to nominate, appoint, re-appoint, remove and replace the General "
       "Manager (and any other manager, director or officer) of the Company, and to determine "
       "the terms of their appointment and remuneration. The Transferees shall vote in favour "
       "of, and shall not oppose, any such nomination, appointment, re-appointment, removal or "
       "replacement.")
clause("7.6", "**Reserved Matters.** Without limiting the generality of the foregoing, none of "
       "the following matters may be approved, undertaken or implemented by the Company, by "
       "any shareholder or by any manager, director or officer of the Company without the "
       "**prior written consent of the Transferor** (acting by the Founder): "
       "(a) any amendment to the MOA, articles or constitutional documents of the Company; "
       "(b) any increase, reduction, reorganisation or other change to the share capital of "
       "the Company; (c) the declaration or payment of any dividend or other distribution to "
       "shareholders; (d) the issuance, allotment, repurchase or redemption of any shares or "
       "other securities; (e) any merger, consolidation, demerger, liquidation, winding-up or "
       "dissolution of the Company; (f) any sale, lease, disposal or acquisition of assets or "
       "businesses outside the ordinary course of business; (g) any incurrence of borrowing, "
       "financial indebtedness, guarantee or surety, or the creation of any Encumbrance over "
       "the assets of the Company; (h) the entry into, amendment or termination of any "
       "related-party or affiliate transaction; (i) the approval of the annual budget, business "
       "plan and audited financial statements; (j) the appointment, removal or remuneration of "
       "the auditors and key management personnel; (k) the commencement, settlement or "
       "compromise of any litigation, arbitration or material claim; (l) any change to the "
       "nature, scope or location of the business activities of the Company; (m) the opening "
       "or closing of any bank account, or any change to authorised signatories thereto; and "
       "(n) any transfer, pledge or grant of any Encumbrance by a shareholder over its shares "
       "in the Company.")
clause("7.7", "**Entrenchment in the Memorandum of Association.** The Parties shall procure "
       "that, at Completion, the Company’s MOA is amended (and notarised before the Notary "
       "Public and registered with the DET) so as to reflect and entrench the arrangements set "
       "out in this Clause 7, including (i) the Transferor’s casting / deciding vote; "
       "(ii) the Transferor’s sole right to appoint and remove the General Manager; and "
       "(iii) the requirement that the Reserved Matters require the prior written consent of "
       "the Transferor. Where, by Applicable Law, any such entrenchment must take the form of "
       "weighted voting rights, super-majority requirements or a dedicated class of shares, the "
       "Parties shall co-operate in good faith to give effect to the same in the manner most "
       "likely to be enforceable under UAE law.")
clause("7.8", "**No Variation Without Founder Consent.** The arrangements in this Clause 7 are "
       "a fundamental condition of the transfer of the Transfer Shares and may not be varied, "
       "suspended or waived save with the prior written consent of the Transferor (acting by "
       "the Founder). Any purported variation or waiver in breach of this Clause shall be null "
       "and void.")
clause("7.9", "**Specific Performance.** The Parties acknowledge that damages would not be an "
       "adequate remedy for any breach of this Clause 7 and that the Transferor shall be "
       "entitled to seek specific performance, injunctive relief and any other equitable "
       "remedy available under Applicable Law.")
clause("7.10", "**Survival.** This Clause 7 shall take effect from Completion and shall continue "
       "in force for so long as the Transferor holds any shares in the Company.")

# 8. Consideration
heading("8", "Consideration")
clause("8.1", "In consideration of the achievement of the Performance Threshold and the past and "
       "continuing performance and services of the Transferees, the Transfer Shares shall be "
       "transferred to the Transferees for an **aggregate nominal consideration of AED 1.00 "
       "(one UAE Dirham)** (or such nominal/par amount as the Notary Public may require to be "
       "stated), the receipt and adequacy of which the Transferor acknowledges.")
clause("8.2", "The Parties confirm that the transfer is by way of a **performance reward and "
       "incentive** and is not a sale at market value. Each Party shall bear its own tax "
       "position arising from the transfer and shall be responsible for obtaining its own "
       "independent tax advice (including in relation to UAE Corporate Tax and any applicable "
       "VAT).")
clause("8.3", "Nothing in this Clause 8 limits the Founder’s discretion under Clause 5 or "
       "the rights reserved to the Transferor under Clause 7, nor obliges the Transferor to "
       "transfer the Transfer Shares otherwise than in accordance with this Agreement.")

# 8. Pre-emption waiver
heading("9", "Waiver of Pre-emption Rights")
clause("9.1", "Each shareholder of the Company (including the Transferees) irrevocably and "
       "unconditionally **waives** any right of pre-emption, right of first refusal, tag-along, "
       "or similar right (whether arising under the CCL, the Company’s MOA or otherwise) to "
       "which it may be entitled in respect of the transfer of the Transfer Shares contemplated "
       "by this Agreement, and consents to such transfer.")
clause("9.2", "Each shareholder shall execute such waivers, consents and resolutions as may be "
       "required to confirm the foregoing for the purposes of notarisation and DET registration.")

# 10. Reps & warranties
heading("10", "Representations and Warranties")
clause("10.1", "Each Party represents and warrants to the others that: (a) it has the power and "
       "authority (and, where applicable, the requisite corporate authorisations) to enter into "
       "and perform this Agreement; (b) this Agreement constitutes its legal, valid and binding "
       "obligations; and (c) its entry into this Agreement does not breach any Applicable Law or "
       "any agreement binding upon it.")
clause("10.2", "The Transferor further represents and warrants that, as at the Effective Date, "
       "it is the legal and beneficial owner of 70% of the issued share capital of the Company, "
       "free from Encumbrances, and that the Transfer Shares will at Completion be transferred "
       "free from Encumbrances.")
clause("10.3", "Each Transferee acknowledges that, save as expressly set out in this Agreement, "
       "no representation, warranty or assurance has been given as to the future profitability "
       "or value of the Company or the Transfer Shares.")

# 11. Restrictions pending completion
heading("11", "Conduct Pending Completion")
clause("11.1", "From the Effective Date until Completion or termination of this Agreement, the "
       "Transferor shall not create any Encumbrance over, or transfer to any third party, the "
       "Transfer Shares in a manner that would prevent it from performing its obligations under "
       "this Agreement, save with the prior written consent of the Founder.")
clause("11.2", "Nothing in this Clause restricts the ordinary conduct of the business of the "
       "Company or the exercise by the Transferor of its rights as a 70% shareholder.")

# 12. Confidentiality
heading("12", "Confidentiality")
clause("12.1", "Each Party shall keep confidential the terms of this Agreement and all "
       "information concerning the other Parties and the Company obtained in connection with it, "
       "and shall not disclose the same save: (a) as required by Applicable Law, a court or "
       "regulator (including the DET and the Notary Public); (b) to its professional advisers "
       "under a duty of confidence; or (c) with the prior written consent of the other Parties.")

# 13. Taxes and costs
heading("13", "Costs, Fees and Taxes")
clause("13.1", "Save as otherwise agreed in writing, the Transferor shall bear the Notary Public "
       "fees, DET amendment fees and registration costs directly arising from Completion, and "
       "each Party shall bear its own legal and advisory costs.")
clause("13.2", "Each Party shall be responsible for its own taxes (if any) arising from the "
       "transfer of the Transfer Shares.")

# 14. Term and termination
heading("14", "Term and Termination")
clause("14.1", "This Agreement takes effect on the Effective Date and continues until the "
       "earlier of (a) Completion and (b) termination in accordance with this Clause 14. The "
       "rights and obligations of the Transferor under Clause 7 shall, however, take effect "
       "from Completion and continue notwithstanding any termination thereafter, for so long as "
       "the Transferor holds any shares in the Company.")
clause("14.2", "This Agreement may be terminated by the written agreement of the Transferor and "
       "the Company, or by the Transferor (acting by the Founder) if the Performance Threshold "
       "has not been achieved and the Transferor reasonably determines that it is no longer "
       "capable of being achieved, in each case without liability save in respect of antecedent "
       "breaches.")
clause("14.3", "Clauses 1, 7, 12, 15 and 16 survive termination.")

# 15. Governing law
heading("15", "Governing Law and Jurisdiction")
clause("15.1", "This Agreement and any non-contractual obligations arising out of or in "
       "connection with it are governed by, and shall be construed in accordance with, the "
       "**federal laws of the United Arab Emirates as applied in the Emirate of Dubai**.")
clause("15.2", "The Parties irrevocably submit to the **exclusive jurisdiction of the Courts of "
       "Dubai** in respect of any dispute arising out of or in connection with this Agreement, "
       "without prejudice to any mandatory requirement for notarisation or registration before "
       "the Notary Public and the DET.")
clause("15.3", "The Parties may, by written agreement, elect to refer any dispute instead to "
       "arbitration administered by the **Dubai International Arbitration Centre (DIAC)**, "
       "seated in Dubai and conducted in the English language.")

# 16. General / Boilerplate
heading("16", "General Provisions")
clause("16.1", "**Entire Agreement.** This Agreement constitutes the entire agreement between "
       "the Parties in relation to its subject matter and supersedes all prior arrangements and "
       "understandings.")
clause("16.2", "**Amendments.** No amendment to this Agreement is effective unless in writing "
       "and signed by or on behalf of each Party; and any amendment to Clause 7 is further "
       "subject to Clause 7.8.")
clause("16.3", "**No Waiver.** No failure or delay in exercising any right operates as a waiver "
       "of it.")
clause("16.4", "**Severability.** If any provision is or becomes invalid or unenforceable, the "
       "remaining provisions continue in full force and the Parties shall negotiate in good "
       "faith to replace the invalid provision with a valid one of similar effect.")
clause("16.5", "**Assignment.** No Party may assign or transfer its rights or obligations under "
       "this Agreement without the prior written consent of the Transferor (acting by the "
       "Founder).")
clause("16.6", "**Further Assurance.** Each Party shall execute such further documents and do "
       "such further acts as may reasonably be required to give full effect to this Agreement, "
       "including the irrevocable proxies and powers of attorney referred to in Clause 7.3.")
clause("16.7", "**Counterparts.** This Agreement may be executed in any number of counterparts, "
       "each of which is an original and all of which together constitute one agreement.")
clause("16.8", "**Notices.** Notices under this Agreement shall be in writing and delivered by "
       "hand, courier or email to the addresses of the Parties set out in this Agreement (or as "
       "notified in writing from time to time).")
clause("16.9", "**Language.** This Agreement is executed in English. Where a notarised Arabic "
       "version is required for registration before the Notary Public and the DET, the Parties "
       "shall procure a certified Arabic translation; in the event of any conflict between the "
       "English and the notarised Arabic version, the **Arabic version shall prevail** for the "
       "purposes of registration before the Dubai authorities, and the English version shall "
       "otherwise govern the commercial intention of the Parties.")

# ----------------------------------------------------------------------------
# EXECUTION
# ----------------------------------------------------------------------------
section_title("Execution")
plain("**IN WITNESS WHEREOF** the Parties have executed this Agreement on the date stated at the "
      "beginning of this Agreement.", space_after=12)


def sig_block(role, entity, name_label, extra_rows=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "FBFAF6")
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "C9C2AE")
        borders.append(e)
    tcPr.append(borders)
    p0 = cell.paragraphs[0]
    no_space(p0)
    add_run(p0, role, size=9, bold=True, color=GOLD, font=HEAD_FONT, caps=True)
    p1 = cell.add_paragraph()
    no_space(p1)
    p1.paragraph_format.space_before = Pt(2)
    add_run(p1, entity, size=11, bold=True, color=NAVY, font=HEAD_FONT)
    rows = (extra_rows or []) + [
        ("Signature", "_______________________________________"),
        (name_label, "_______________________________________"),
        ("Title / Capacity", "_______________________________________"),
        ("Date", "_______________________________________"),
    ]
    for label, line in rows:
        rp = cell.add_paragraph()
        no_space(rp)
        rp.paragraph_format.space_before = Pt(7)
        add_run(rp, f"{label}:  ", size=9.5, bold=True, color=GREY)
        add_run(rp, line, size=9.5, color=BLACK)
    pad = cell.add_paragraph()
    no_space(pad)
    add_run(pad, "", size=4)
    doc.add_paragraph()


sig_block("The Transferor", "PODOCARPUS REAL ESTATE LLC",
          "Name of Authorised Signatory")
sig_block("Acting by the Founder", "FOR AND ON BEHALF OF THE FOUNDER",
          "Name of Founder")
sig_block("The Company", "ORIZURU HOLIDAY HOMES LLC",
          "Name of Authorised Signatory")
sig_block("Transferee (1)", "MR. YSLAM ATAYEV", "Name")
sig_block("Transferee (2)", "MR. ORAZDURDY MAMIYEV", "Name")

plain("Witnessed by: ______________________________   Name: ______________________________   "
      "Date: ____________", space_after=4, size=9.5)

# ----------------------------------------------------------------------------
# SCHEDULES
# ----------------------------------------------------------------------------
doc.add_page_break()
section_title("Schedule 1 — Existing Shareholding (as at the Effective Date)")


def make_table(headers, rows, widths=None, total_row=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "0E2A47")
        hp = hdr[i].paragraphs[0]
        no_space(hp)
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(hp, h, size=9.5, bold=True, color=WHITE, font=HEAD_FONT)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        shade = "F4F1E8" if r_idx % 2 else "FFFFFF"
        for i, val in enumerate(row):
            set_cell_bg(cells[i], shade)
            cp = cells[i].paragraphs[0]
            no_space(cp)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(cp, val, size=9.5)
    if total_row:
        cells = t.add_row().cells
        for i, val in enumerate(total_row):
            set_cell_bg(cells[i], "E7DFC8")
            cp = cells[i].paragraphs[0]
            no_space(cp)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(cp, val, size=9.5, bold=True, color=NAVY)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


make_table(
    ["Shareholder", "Shareholding %", "Role"],
    [
        ["Podocarpus Real Estate LLC", "70%", "Majority shareholder / Transferor"],
        ["Mr. Yslam Atayev", "16%", "Minority shareholder / Transferee (1)"],
        ["Mr. Orazdurdy Mamiyev", "14%", "Minority shareholder / Transferee (2)"],
    ],
    widths=[2.6, 1.4, 2.6],
    total_row=["Total", "100%", ""],
)

# Schedule 2
section_title("Schedule 2 — Determination of the Performance Threshold")
plain("**1. Definition of Cumulative Gross Profit.** “Cumulative Gross Profit” means, as "
      "at any date of measurement, the aggregate of the Company’s gross profit accumulated "
      "from the commencement of the Company’s operations to that date, where gross profit for "
      "any period equals **total revenue actually realised and received by the Company** in "
      "respect of its holiday-homes / short-term rental and related activities, **less the direct "
      "cost of generating that revenue** (including owner payouts, OTA / platform commissions, "
      "cleaning, maintenance, utilities directly attributable to operated units, and similar "
      "direct costs).")
plain("**2. Threshold composition.** The Performance Threshold is satisfied when Cumulative Gross "
      "Profit first equals or exceeds **AED 400,000**, deemed to comprise:")
make_table(
    ["Component", "Amount (AED)", "Purpose"],
    [
        ["Capital Recovery Amount", "200,000", "Full recovery of the Initial Capital of AED 200,000"],
        ["Surplus Profit Amount", "200,000", "Additional profit generated for the Company"],
    ],
    widths=[2.3, 1.6, 2.7],
    total_row=["Performance Threshold", "400,000", "Trigger Event arises"],
)
plain("**3. Measurement.** Cumulative Gross Profit shall be determined from the Company’s "
      "books and records and certified by the Company’s accountant or auditor in the "
      "Threshold Certificate. The Founder may require independent verification under Clause 3.4. "
      "There is **no fixed deadline** for achieving the Performance Threshold.", space_after=4)
plain("**4. Adjustments.** Extraordinary, non-recurring or non-operating items may be excluded "
      "from the calculation at the reasonable determination of the Founder, acting in good faith, "
      "to ensure the figure reflects genuine operating performance.", space_after=4)

# Schedule 3
section_title("Schedule 3 — Form of Allocation Notice")
plain("[On the letterhead of Podocarpus Real Estate LLC]", italic=True, color=GREY, space_after=8)
plain("**ALLOCATION NOTICE**", justify=False, space_after=6)
plain("Date: ____ / ____ / 20____", justify=False, space_after=4)
plain("To: Mr. Yslam Atayev and Mr. Orazdurdy Mamiyev, and to Orizuru Holiday Homes LLC.",
      space_after=8)
plain("We refer to the Share Transfer Agreement dated ____ / ____ / 20____ (the “Agreement”). "
      "Capitalised terms used in this notice have the meanings given in the Agreement.", space_after=6)
plain("We confirm that the Trigger Event has occurred. Pursuant to Clauses 4 and 5 of the "
      "Agreement, the Founder has determined, in the Founder’s absolute and sole discretion "
      "and according to the respective performance of the Transferees, that the Transfer Shares "
      "(aggregating 21%) shall be allocated as follows:", space_after=8)
make_table(
    ["Transferee", "% of Company allocated (out of 21%)"],
    [
        ["Mr. Yslam Atayev", "________ %"],
        ["Mr. Orazdurdy Mamiyev", "________ %"],
    ],
    widths=[3.4, 3.0],
    total_row=["Total transferred", "21%"],
)
plain("This Allocation Notice is final and binding in accordance with Clause 4.3 of the "
      "Agreement. The Parties shall proceed to Completion in accordance with Clause 6.",
      space_after=10)
plain("Signed for and on behalf of the Founder / Podocarpus Real Estate LLC:", space_after=10)
plain("____________________________________        Name: ____________________________",
      justify=False, space_after=4)

# Schedule 4
section_title("Schedule 4 — Illustrative Resulting Shareholding (Non-Binding)")
plain("The following scenarios are illustrative only and do not bind or fetter the Founder’s "
      "discretion under Clause 4. In every case Podocarpus moves from 70% to 49% (a reduction of "
      "21%).", space_after=8)
plain("**Scenario A — Split proportional to existing holdings (16:14):**", justify=False, space_after=4)
make_table(
    ["Shareholder", "Before", "Transfer", "After"],
    [
        ["Podocarpus Real Estate LLC", "70%", "− 21%", "49.0%"],
        ["Mr. Yslam Atayev", "16%", "+ 11.2%", "27.2%"],
        ["Mr. Orazdurdy Mamiyev", "14%", "+ 9.8%", "23.8%"],
    ],
    widths=[2.8, 1.2, 1.3, 1.3],
    total_row=["Total", "100%", "", "100%"],
)
plain("**Scenario B — Equal split (10.5% each):**", justify=False, space_after=4)
make_table(
    ["Shareholder", "Before", "Transfer", "After"],
    [
        ["Podocarpus Real Estate LLC", "70%", "− 21%", "49.0%"],
        ["Mr. Yslam Atayev", "16%", "+ 10.5%", "26.5%"],
        ["Mr. Orazdurdy Mamiyev", "14%", "+ 10.5%", "24.5%"],
    ],
    widths=[2.8, 1.2, 1.3, 1.3],
    total_row=["Total", "100%", "", "100%"],
)
plain("The actual allocation shall be as set out in the Allocation Notice (Schedule 3), "
      "determined by the Founder according to each Transferee’s respective performance.",
      space_after=4)

# Schedule 5 — items to complete
section_title("Schedule 5 — Information to be Completed Before Execution")
plain("The following items, shown as [ ● ] or blank lines in this Agreement, must be "
      "completed by the Parties (with their legal adviser) before signing and notarisation:",
      space_after=8)
for item in [
    "Effective Date of the Agreement.",
    "Podocarpus Real Estate LLC — trade licence number and registered address.",
    "Orizuru Holiday Homes LLC — trade licence number and registered address.",
    "Full legal name of the Founder of Podocarpus, and passport / Emirates ID number.",
    "Mr. Yslam Atayev — passport / Emirates ID number, nationality and address.",
    "Mr. Orazdurdy Mamiyev — passport / Emirates ID number, nationality and address.",
    "Confirmation of the Company’s par value per share and total number of shares (for the notarised transfer).",
    "Names and titles of the authorised signatories for each company.",
    "Whether disputes are to be resolved by Dubai Courts (Clause 15.2) or DIAC arbitration (Clause 15.3).",
    "Whether to add a specific AED threshold to Reserved Matter 7.6(f) (asset disposals outside ordinary course).",
    "Confirmation that the amended MOA will entrench the Transferor's casting vote, Reserved Matters and sole right to appoint the General Manager (Clause 7.7).",
    "Certified Arabic translation for notarisation and DET registration.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, item, size=10)

# Closing note
doc.add_paragraph()
callout(
    "Reminder",
    ["This is a binding commercial template. To be legally effective in Dubai, the share "
     "transfer must be notarised before a Dubai Notary Public and registered with the DET, "
     "with the MOA amended accordingly. Have a UAE-licensed lawyer review and finalise this "
     "document and prepare the Arabic notarised addendum prior to execution."],
)

import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, "Orizuru_Share_Transfer_Agreement.docx")
doc.save(out_path)
print("Saved:", out_path)
